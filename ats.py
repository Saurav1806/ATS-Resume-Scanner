import os
import io
import base64
from docx import Document
import google.generativeai as genai
import pdf2image

#Add Your GenAI api key
genai.configure(api_key="YOUR_API_KEY_HERE")

def process_pdf(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        images = pdf2image.convert_from_path(file_path, first_page=1, last_page=1)
        first_page = images[0]
        img_byte_arr = io.BytesIO()
        first_page.save(img_byte_arr, format='JPEG')
        img_byte_arr = img_byte_arr.getvalue()
        pdf_parts = [
            {
                "mime_type": "image/jpeg",
                "data": base64.b64encode(img_byte_arr).decode()
            }
        ]
        return pdf_parts
    except Exception as e:
        raise RuntimeError(f"Error processing PDF: {e}")

def analyze_resume(input_text, pdf_content, prompt):
    """Sends resume and job description to Gemini AI for analysis."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([input_text, pdf_content[0], prompt])
    return response.text

def save_to_word(file_path, content):
    """Saves the output analysis to a Word document."""
    doc = Document()
    doc.add_heading("Resume Analysis", level=1)
    doc.add_paragraph(content)
    output_path = os.path.splitext(file_path)[0] + "_output.docx"
    doc.save(output_path)
    return output_path

def main():
    # Resume File Path
    path = 'C:\\Documents\\ATS\\'
    file = input("Enter the full path to the resume file (PDF): ").strip() #File Name as Input
    file_path = os.path.join(path, file)
    print(f"File path: {file_path}")

    if not os.path.isfile(file_path):
        print("Error: File not found. Please check the path and try again.")
        return

    # Input: Job Description
    job_description = input("Enter the job description: ").strip()

    # Prompts
    prompt_analysis = (
        "You are an experienced Technical Human Resource Manager. Analyze the given resume "
        "against the provided job description. Highlight strengths, weaknesses, missing skills, "
        "and overall fit. Suggest improvements and provide a detailed ATS score and acceptance percentage."
    )

    prompt_missing_skills = (
        "Analyze the resume and job description deeply. List all the skills, certifications, or experiences "
        "mentioned in the job description but missing from the resume. Suggest how the candidate can address these gaps."
    )

    prompt_improvement_suggestions = (
        "Evaluate the resume and provide actionable suggestions for improvement. Focus on formatting, content clarity, "
        "keyword optimization, and relevance to the job description. Provide a structured list of suggestions."
    )

    prompt_acceptance_probability = (
        f"Based on the resume and {job_description} as job description, estimate the likelihood (in percentage) of the resume being accepted "
        "by an ATS system. Justify your reasoning with detailed feedback. Include strengths, weaknesses, and missing elements."
    )

    # Process PDF and Analyze
    try:
        print("Processing resume...\n")
        pdf_content = process_pdf(file_path)

        print("Analyzing resume...\n")
        analysis = analyze_resume(job_description, pdf_content, prompt_analysis)

        print("Identifying missing skills...\n")
        missing_skills = analyze_resume(job_description, pdf_content, prompt_missing_skills)

        print("Generating improvement suggestions...\n")
        improvement_suggestions = analyze_resume(job_description, pdf_content, prompt_improvement_suggestions)

        print("Calculating acceptance probability...\n")
        acceptance_probability = analyze_resume(job_description, pdf_content, prompt_acceptance_probability)

        print("Saving analysis to Word document...\n")
        output_content = f"""
        Detailed Resume Analysis:

        1. Overall Analysis: \n
        {analysis}

        2. Missing Skills and Certifications: \n
        {missing_skills}

        3. Improvement Suggestions: \n
        {improvement_suggestions}

        4. ATS Acceptance Probability: \n
        {acceptance_probability}
        """

        output_path = save_to_word(file_path, output_content)

        print(f"Analysis saved successfully to: {output_path}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()
